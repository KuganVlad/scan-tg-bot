from aiogram.dispatcher.filters.state import State, StatesGroup
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram import Bot, Dispatcher, types, executor
from aiogram.dispatcher import FSMContext

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import docx.oxml.ns as ns
from docx import Document

from apscheduler.schedulers.asyncio import AsyncIOScheduler
from dateutil import parser as date_parser
import configparser
import subprocess
import reports
import sqlite3
import logging
import locale
import os

logging.basicConfig(filename='bot.log', level=logging.ERROR,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')

locale.setlocale(locale.LC_TIME, 'ru_RU.UTF-8')
conn = sqlite3.connect('data.db')
cursor = conn.cursor()

# Считываем учетные данные, запускаем клиент
config = configparser.ConfigParser()
config.read("config.ini")
TOKEN = config['Telegram']['bot_token']

storage = MemoryStorage()
bot = Bot(token=f'{TOKEN}')
dp = Dispatcher(bot, storage=storage)


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    page_num_run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)


class YourState(StatesGroup):
    wait_count_pop_news = State()
    wait_count_key_word = State()
    wait_key_word = State()


def clear_database():
    conn = sqlite3.connect('data.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM news")
    conn.commit()
    conn.close()


async def execute_code():
    subprocess.run(["python", "main.py"], capture_output=True, text=True)


scheduler = AsyncIOScheduler()
scheduler.add_job(clear_database, 'cron', hour=6, minute=59)
scheduler.add_job(execute_code, 'cron', hour='*', minute=0)
scheduler.start()


def create_tables():
    connection = sqlite3.connect("data.db")
    cursor = connection.cursor()

    cursor.execute('''CREATE TABLE IF NOT EXISTS key_words (
                        word_id INTEGER PRIMARY KEY AUTOINCREMENT,
                        word TEXT,
                        user_id TEXT
                    );''')

    cursor.execute('''CREATE TABLE IF NOT EXISTS news (
                        news_id INTEGER PRIMARY KEY AUTOINCREMENT,
                        dialog_name TEXT,
                        dialog_type INTEGER,
                        dialog_id INTEGER,
                        message_id INTEGER,
                        publication_date TEXT,
                        message_text TEXT,
                        message_media INTEGER,
                        sender_id INTEGER
                    );''')

    cursor.execute('''CREATE TABLE IF NOT EXISTS settings_bot (
                        setting_id INTEGER PRIMARY KEY AUTOINCREMENT,
                        count_pop_news TEXT,
                        count_keyword_news TEXT,
                        arr_keyword TEXT,
                        setting_user_id INTEGER
                    );''')

    cursor.execute('''CREATE TABLE IF NOT EXISTS users (
                        user_id TEXT
                    );''')

    connection.commit()
    connection.close()

def get_allowed_users():
    cursor.execute("SELECT user_id FROM users")
    result = cursor.fetchall()
    return [int(row[0]) for row in result]


def is_user_allowed(user_id):
    allowed_users = get_allowed_users()
    return user_id in allowed_users


def get_total_chats():
    cursor.execute("SELECT COUNT(DISTINCT dialog_name) FROM news")
    result = cursor.fetchone()
    return result[0]


def get_chat_names():
    cursor.execute("SELECT DISTINCT dialog_name FROM news")
    result = cursor.fetchall()
    return [row[0] for row in result]


def get_total_messages():
    cursor.execute("SELECT COUNT(*) FROM news")
    result = cursor.fetchone()
    return result[0]


def get_admin_messages():
    cursor.execute("SELECT COUNT(*) FROM news WHERE sender_id IS NULL")
    result = cursor.fetchone()
    return result[0]


def get_user_messages():
    cursor.execute("SELECT COUNT(*) FROM news WHERE sender_id IS NOT NULL")
    result = cursor.fetchone()
    return result[0]


def get_unique_messages():
    cursor.execute("SELECT COUNT(DISTINCT message_text) FROM news")
    result = cursor.fetchone()
    return result[0]


def get_popular_news_comment(user_id):
    query = f"SELECT message_text, dialog_name, publication_date, dialog_type FROM news WHERE message_text IS NOT NULL ORDER BY dialog_type DESC LIMIT {str(get_setting_count_pop_news(user_id))[2:-2]}"
    cursor.execute(query)
    result = cursor.fetchall()
    news_list = []
    for row in result:
        message_text = row[0]
        if not message_text:
            continue
        # TODO: replace an array with ignored keywords with a table in the database
        elif any(substring in message_text for substring in ["Предлагаем вакансии в Англии."]):
            continue
        dialog_name = row[1]
        publication_date_str = row[2]
        count_comments = row[3]
        if not publication_date_str:
            continue
        publication_date = date_parser.parse(publication_date_str)
        news_list.append((message_text, dialog_name, publication_date, count_comments))
    return news_list


def get_popular_news_comment_for_doc(user_id):
    query = f"SELECT message_text, dialog_name, publication_date, dialog_type, dialog_id, message_id FROM news WHERE message_text IS NOT NULL ORDER BY dialog_type DESC LIMIT {str(get_setting_count_pop_news(user_id))[2:-2]}"
    cursor.execute(query)
    result = cursor.fetchall()
    news_list = []
    for row in result:
        message_text = row[0]
        if not message_text:
            continue
        # TODO: replace an array with ignored keywords with a table in the database
        elif any(substring in message_text for substring in ["Предлагаем вакансии в Англии."]):
            continue
        dialog_name = row[1]
        publication_date_str = row[2]
        count_comments = row[3]
        dialog_id = row[4]
        message_id = row[5]
        if not publication_date_str:
            continue
        publication_date = date_parser.parse(publication_date_str)
        news_list.append((message_text, dialog_name, publication_date, count_comments, dialog_id, message_id))
    return news_list


def get_popular_news(user_id):
    query = f"SELECT message_text, dialog_name, publication_date, COUNT(*) as count FROM news WHERE message_text IS NOT NULL GROUP BY message_text ORDER BY count DESC LIMIT {str(get_setting_count_pop_news(user_id))[2:-2]}"
    cursor.execute(query)
    result = cursor.fetchall()
    news_list = []
    for row in result:
        message_text = row[0]
        if not message_text:
            continue
        elif any(
                # TODO: replace an array with ignored keywords with a table in the database
                substring in message_text for substring in ["Предлагаем вакансии в Англии."]):
            continue
        dialog_name = row[1]
        publication_date_str = row[2]
        if not publication_date_str:
            continue
        publication_date = date_parser.parse(publication_date_str)
        news_list.append((message_text, dialog_name, publication_date))
    return news_list


def get_popular_news_for_doc(user_id):
    query = f"SELECT message_text, dialog_name, publication_date, dialog_id, message_id, COUNT(*) as count FROM news WHERE message_text IS NOT NULL GROUP BY message_text ORDER BY count DESC LIMIT {str(get_setting_count_pop_news(user_id))[2:-2]}"
    cursor.execute(query)
    result = cursor.fetchall()
    news_list = []
    for row in result:
        message_text = row[0]
        if not message_text:
            continue
        elif any(
                # TODO: replace an array with ignored keywords with a table in the database
                substring in message_text for substring in ["Предлагаем вакансии в Англии."]):
            continue
        dialog_name = row[1]
        publication_date_str = row[2]
        dialog_id = row[3]
        message_id = row[4]
        if not publication_date_str:
            continue
        publication_date = date_parser.parse(publication_date_str)
        news_list.append((message_text, dialog_name, publication_date, dialog_id, message_id))
    return news_list


def get_news_by_keywords(user_id):
    keywords = get_setting_arr_keyword(user_id)
    count_limit = str(get_setting_count_key_word_news(user_id))[2:-2]

    keywords = [keyword.strip() for keyword in keywords[0].split(',')]

    query = "SELECT message_text, dialog_name, publication_date, sender_id, COUNT(*) as count FROM news WHERE "
    conditions = []
    values = []
    for keyword in keywords:
        conditions.append("message_text LIKE ?")
        values.append(f"%{keyword}%")
    query += " OR ".join(conditions)
    query += " GROUP BY message_text ORDER BY count DESC LIMIT ?"
    values.append(count_limit)

    cursor.execute(query, values)
    result = cursor.fetchall()
    news_list = []
    for row in result:
        message_text = row[0]
        dialog_name = row[1]
        publication_date_str = row[2]
        publication_date = date_parser.parse(publication_date_str)
        count = row[3]
        news_list.append((message_text, dialog_name, publication_date, count))

    return news_list


def get_news_by_keywords_for_doc(user_id):
    keywords = get_setting_arr_keyword(user_id)
    count_limit = str(get_setting_count_key_word_news(user_id))[2:-2]

    keywords = [keyword.strip() for keyword in keywords[0].split(',')]

    query = "SELECT message_text, dialog_name, publication_date, sender_id, dialog_id, message_id,COUNT(*) as count FROM news WHERE "
    conditions = []
    values = []
    for keyword in keywords:
        conditions.append("message_text LIKE ?")
        values.append(f"%{keyword}%")
    query += " OR ".join(conditions)
    query += " GROUP BY message_text ORDER BY count DESC LIMIT ?"
    values.append(count_limit)

    cursor.execute(query, values)
    result = cursor.fetchall()
    news_list = []
    for row in result:
        message_text = row[0]
        dialog_name = row[1]
        publication_date_str = row[2]
        publication_date = date_parser.parse(publication_date_str)
        count = row[3]
        dialog_id = row[4]
        message_id = row[5]
        news_list.append((message_text, dialog_name, publication_date, count, dialog_id, message_id))

    return news_list


def get_message_text_for_report_ods(dialog_type, keywords):
    query = "SELECT message_text, dialog_name, publication_date, dialog_id, message_id FROM news WHERE dialog_id = ? AND ("
    conditions = []
    values = [dialog_type]
    for keyword in keywords:
        conditions.append("message_text LIKE ?")
        values.append(f"%{keyword}%")
    query += " OR ".join(conditions)
    query += ")"

    cursor.execute(query, values)
    result = cursor.fetchall()

    cursor.close()
    conn.close()

    news_list = []
    for row in result:
        message_text = row[0]
        dialog_name = row[1]
        publication_date_str = row[2]
        publication_date = date_parser.parse(publication_date_str)
        dialog_id = row[3]
        message_id = row[4]
        news_list.append((message_text, dialog_name, publication_date, dialog_id, message_id))
    return news_list


def get_setting_count_pop_news(user_id):
    query = f"SELECT count_pop_news FROM settings_bot WHERE setting_user_id = {user_id}"
    cursor.execute(query)
    result = cursor.fetchall()
    return [row[0] for row in result]


def get_setting_count_key_word_news(user_id):
    query = f"SELECT count_keyword_news FROM settings_bot WHERE setting_user_id = {user_id}"
    cursor.execute(query)
    result = cursor.fetchall()
    return [row[0] for row in result]


def get_setting_arr_keyword(user_id):
    query = f"SELECT arr_keyword FROM settings_bot WHERE setting_user_id = {user_id}"
    cursor.execute(query)
    result = cursor.fetchall()
    return [row[0] for row in result]


def set_setting_count_pop_news(user_id, count):
    cursor.execute('UPDATE settings_bot SET count_pop_news = ? WHERE setting_user_id = ?', (count, user_id))
    conn.commit()


def set_setting_count_key_word_news(user_id, count):
    cursor.execute('UPDATE settings_bot SET count_keyword_news = ? WHERE setting_user_id = ?', (count, user_id))
    conn.commit()


def set_setting_arr_keyword(user_id, arr_word):
    cursor.execute('UPDATE settings_bot SET arr_keyword = ? WHERE setting_user_id = ?', (",".join(arr_word), user_id))
    conn.commit()


# Функция начала работы
@dp.message_handler(commands=['start'])
async def start_question(message: types.Message):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        buttons = [
            "Новости",
            "Цифровые данные",
            "Отчёты"
        ]
        keyboard.add(*buttons)
        await message.answer(
            "Привет! Данный бот анализирует сообщения за текущий день.\nДля продолжения работы выберите интересующий раздел или перейдите в раздел помощи (/help)",
            reply_markup=keyboard)
    else:
        await message.answer(
            "Доступ запрещён. Для получения дополнительной информации перейдите в раздел помощи (/help).")


# Функция помощи
@dp.message_handler(commands=['help'])
async def help_question(message: types.Message):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        buttons = ["Вернуться в главное меню"]
        keyboard.add(*buttons)
        await message.answer("Данный бот предоставляет информацию о сообщениях, размещённых в каналах и чатах.\n"
                             "Сообщения обновляются раз в час и ежедневно удаляются из базы в 07:00.\n"
                             "При возникновении каких-либо проблем или наличии предложений свяжитесь с администратором",
                             reply_markup=keyboard)
    else:
        await message.answer("Для получения доступа к боту свяжитесь с администратором.")


# Функция очистки базы данных
@dp.message_handler(commands=['clear_db'])
async def clear_database_handler(message: types.Message):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        clear_database()
        keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        buttons = ["Вернуться в главное меню"]
        keyboard.add(*buttons)
        await message.answer("База данных очищена", reply_markup=keyboard)
    else:
        await message.answer("Для получения доступа к боту свяжитесь с администратором.")


# Функция обновления данных
@dp.message_handler(commands=['get_data'])
async def get_data_handler(message: types.Message):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        await message.answer("Обновление данных займёт от 5 до 10 минут")
        await execute_code()
        await message.answer("Данные обновлены")
    else:
        # Сообщение о закрытом доступе
        await message.answer("Для получения доступа к боту свяжитесь с администратором.")


# Функция доступа к настройкам
@dp.message_handler(commands=['settings'])
async def settings_handler(message: types.Message):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        buttons = ["Количество популярных новостей", "Количество новостей по ключ. словам",
                   "Список ключевых слов", "Вернуться в главное меню"]
        keyboard.add(*buttons)
        await message.answer("Изменение настроек будет применено только для вашего профиля", reply_markup=keyboard)
    else:
        # Сообщение о закрытом доступе
        await message.answer("Для получения доступа к боту свяжитесь с администратором.")


# Функция возврата в главное меню
@dp.message_handler(lambda message: message.text == "Вернуться в главное меню")
async def return_start(message: types.Message):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
        buttons = [
            "Новости",
            "Цифровые данные",
            "Отчёты"
        ]
        keyboard.add(*buttons)
        await message.answer("Выбери интересующий раздел", reply_markup=keyboard)
    else:
        await message.answer("Доступ закрыт.")


@dp.message_handler()
async def handle_button_click(message: types.Message):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        if message.text == "Количество анализируемых чатов и каналов":
            total_chats = get_total_chats()
            await message.answer(f"Общее количество анализируемых чатов и каналов: {total_chats}")
        elif message.text == "Название анализируемых чатов и каналов":
            chat_names = get_chat_names()
            formatted_names = "\n".join(chat_names)
            await message.answer(f"Название анализируемых чатов и каналов:\n{formatted_names}")
        elif message.text == "Общее количество сообщений":
            total_messages = get_total_messages()
            await message.answer(f"Общее количество сообщений в чатах и каналах: {total_messages}")
        elif message.text == "Количество сообщений администраторов":
            admin_messages = get_admin_messages()
            await message.answer(f"Количество сообщений администраторов: {admin_messages}")
        elif message.text == "Количество сообщений пользователей":
            user_messages = get_user_messages()
            await message.answer(f"Количество сообщений пользователей: {user_messages}")
        elif message.text == "Количество уникальных сообщений":
            unique_messages = get_unique_messages()
            await message.answer(f"Количество уникальных сообщений: {unique_messages}")
        elif message.text == "Новости по комментариям":
            popular_news = get_popular_news_comment(user_id)

            if popular_news:
                doc = Document()

                for index, news in enumerate(popular_news, 1):
                    message_text = news[0]
                    dialog_name = news[1]
                    publication_date = news[2]
                    count_comments = news[3]
                    formatted_date = publication_date.strftime("%H:%M:%S %d:%m:%Y")
                    # Добавляем информацию о новости в документ
                    index_dia_name_paragraph = doc.add_paragraph()
                    index_dia_name_run = index_dia_name_paragraph.add_run(f"{index}) {dialog_name}")
                    index_dia_name_run.bold = True
                    doc.add_paragraph(f"{formatted_date}")
                    doc.add_paragraph(f"{message_text}")
                    doc.add_paragraph(f"Количество комментариев: {count_comments}")
                    doc.add_paragraph()

                # Сохраняем документ в файл
                doc.save("Popular_news_comment.docx")

                # Отправляем файл через бот пользователю
                with open("Popular_news_comment.docx", "rb") as file:
                    await message.answer_document(file)

                # Удаляем файл с устройства
                os.remove("Popular_news_comment.docx")

            else:
                await message.answer("Нет популярных новостей")
        elif message.text == "Популярные новости":
            popular_news = get_popular_news(user_id)

            if popular_news:
                doc = Document()

                for index, news in enumerate(popular_news, 1):
                    message_text = news[0]
                    dialog_name = news[1]
                    publication_date = news[2]
                    formatted_date = publication_date.strftime("%H:%M:%S %d:%m:%Y")
                    # Добавляем информацию о новости в документ
                    index_dia_name_paragraph = doc.add_paragraph()
                    index_dia_name_run = index_dia_name_paragraph.add_run(f"{index}) {dialog_name}")
                    index_dia_name_run.bold = True
                    doc.add_paragraph(f"{formatted_date}")
                    doc.add_paragraph(f"{message_text}")
                    doc.add_paragraph()

                # Сохраняем документ в файл
                doc.save("Popular_news.docx")

                # Отправляем файл через бот пользователю
                with open("Popular_news.docx", "rb") as file:
                    await message.answer_document(file)

                # Удаляем файл с устройства
                os.remove("Popular_news.docx")

            else:
                await message.answer("Нет популярных новостей")
        elif message.text == "Новости по ключевым словам":
            try:
                news_by_keywords = get_news_by_keywords(user_id)
                if news_by_keywords:
                    doc = Document()

                    for index, news in enumerate(news_by_keywords, 1):
                        message_text = news[0]
                        dialog_name = news[1]
                        publication_date = news[2]
                        if news[3] != None:
                            sender = news[3]
                        else:
                            sender = "администратора канала(группы)"
                        formatted_date = publication_date.strftime("%H:%M:%S %d:%m:%Y")

                        # Добавляем информацию о новости в документ
                        index_dia_name_paragraph = doc.add_paragraph()
                        index_dia_name_run = index_dia_name_paragraph.add_run(f"{index}) {dialog_name}")
                        index_dia_name_run.bold = True
                        doc.add_paragraph(f"{formatted_date}")
                        doc.add_paragraph(f"{message_text}")
                        doc.add_paragraph(f"Опубликовано от : {sender}")
                        doc.add_paragraph()

                    # Сохраняем документ в файл
                    doc.save("New_key_word.docx")

                    # Отправляем файл через бот пользователю
                    with open("New_key_word.docx", "rb") as file:
                        await message.answer_document(file)

                    # Удаляем файл с устройства
                    os.remove("New_key_word.docx")

                else:
                    await message.answer("Нет новостей, содержащих ключевые слова")
            except sqlite3.OperationalError:
                await message.answer("Ключевые слова отсутствую")
        elif message.text == "Количество популярных новостей":
            user_id = message.from_user.id
            if is_user_allowed(user_id):
                num = str(get_setting_count_pop_news(user_id))[2:-2]
                # Запуск состояния ожидания вопроса
                await YourState.wait_count_pop_news.set()
                await message.answer(
                    f"Сейчас в выборку попадает {num} популярных новостей. "
                    f"\nДля изменения количества введи новое число: ")
            else:
                # Сообщение о закрытом доступе
                await message.answer("Доступ закрыт.")
        elif message.text == "Количество новостей по ключ. словам":
            user_id = message.from_user.id
            if is_user_allowed(user_id):
                num = str(get_setting_count_key_word_news(user_id))[2:-2]
                # Запуск состояния ожидания вопроса
                await YourState.wait_count_key_word.set()
                await message.answer(
                    f"Сейчас в выборку попадает {num} новостей по ключевым словам. "
                    f"\nДля изменения количества введи новое число: ")
            else:
                # Сообщение о закрытом доступе
                await message.answer("Доступ закрыт.")
        elif message.text == "Список ключевых слов":
            user_id = message.from_user.id
            if is_user_allowed(user_id):
                arr_word = str(get_setting_arr_keyword(user_id))[2:-2]
                # Запуск состояния ожидания вопроса
                await YourState.wait_key_word.set()
                await message.answer(
                    f"Сейчас поиск происходи по следующим словам: \n{arr_word}. "
                    f"\nДля их изменения введи новый список через , ")
            else:
                # Сообщение о закрытом доступе
                await message.answer("Доступ закрыт.")
        elif message.text == "Новости" or message.text == "Вернуться к новостям":
            user_id = message.from_user.id
            if is_user_allowed(user_id):
                keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
                buttons = [
                    "Популярные новости",
                    "Новости по ключевым словам",
                    "Новости по комментариям",
                    "Вернуться в главное меню"
                ]
                keyboard.add(*buttons)
                await message.answer(
                    "Выбери интересующий раздел",
                    reply_markup=keyboard)
            else:
                # Сообщение о закрытом доступе
                await message.answer(
                    "Доступ закрыт.")
        elif message.text == "Цифровые данные":
            user_id = message.from_user.id
            if is_user_allowed(user_id):
                keyboard = types.ReplyKeyboardMarkup(resize_keyboard=True)
                buttons = [
                    "Название анализируемых чатов и каналов",
                    "Количество анализируемых чатов и каналов",
                    "Общее количество сообщений",
                    "Количество сообщений администраторов",
                    "Количество сообщений пользователей",
                    "Количество уникальных сообщений",
                    "Вернуться в главное меню"
                ]
                keyboard.add(*buttons)
                await message.answer(
                    "Выбери интересующий раздел",
                    reply_markup=keyboard)
            else:
                # Сообщение о закрытом доступе
                await message.answer(
                    "Доступ закрыт.")
        elif message.text == "Отчёты":
            await reports.handle_reports(message)
        elif message.text == "Сведения по распоряжению":
            await reports.handle_reports_rasp(message)
        elif message.text == "Сведения для руководства":
            await reports.handle_reports_manager(message)
        elif message.text == "По комментариям":
            await reports.handle_manager_comments(message, user_id)
        elif message.text == "По ключевым словам":
            await reports.handle_manager_keyword(message, user_id)
        elif message.text == "По популярным новостям":
            await reports.handle_manager_popular(message, user_id)
        else:
            await message.answer("Неизвестный метод")
    else:
        await message.answer("Для получения доступа к боту свяжитесь с администратором.")


# Функция для обработки ввода количества популярных новостей
@dp.message_handler(state=YourState.wait_count_pop_news)
async def process_count_pop_news(message: types.Message, state: FSMContext):
    try:
        count = int(message.text)
        if count > 0:
            user_id = message.from_user.id
            if is_user_allowed(user_id):
                set_setting_count_pop_news(user_id, count)  # Записываем значение в базу данных
                await message.answer(f"Количество популярных новостей успешно обновлено.")
            else:
                await message.answer("Доступ закрыт.")
        else:
            await message.answer("Следующий раз, введите число, больше нуля.")
    except ValueError:
        await message.answer("Неверный формат числа. Следующий раз, введите целое число.")

    await state.finish()


# Функция для обработки ввода количества новостей по ключевым словам
@dp.message_handler(state=YourState.wait_count_key_word)
async def process_count_key_word_news(message: types.Message, state: FSMContext):
    try:
        count = int(message.text)
        if count > 0:
            user_id = message.from_user.id
            if is_user_allowed(user_id):
                set_setting_count_key_word_news(user_id, count)  # Записываем значение в базу данных
                await message.answer(f"Количество новостей по ключевым словам успешно обновлено.")
            else:
                await message.answer("Доступ закрыт.")
        else:
            await message.answer("Следующий раз, введите число, больше нуля.")
    except ValueError:
        await message.answer("Неверный формат числа. Следующий раз, введите целое число.")

    await state.finish()


# Функция для обработки ввода списка ключевых слов
@dp.message_handler(state=YourState.wait_key_word)
async def process_key_word(message: types.Message, state: FSMContext):
    user_id = message.from_user.id
    if is_user_allowed(user_id):
        arr_word = message.text.split(",")
        set_setting_arr_keyword(user_id, arr_word)
        await message.answer(f"Список ключевых слов успешно обновлен.")
    else:
        await message.answer("Доступ закрыт.")

    await state.finish()


if __name__ == '__main__':
    while True:
        try:
            create_tables()
            executor.start_polling(dp, timeout=60, skip_updates=True)
        except Exception as e:
            print(f"Error occurred: {e}")
            print("Restarting the bot...")
